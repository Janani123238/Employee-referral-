"""
Real resume text extraction — no mocks.

Supports:
  - .pdf   -> native text layer via pdfplumber, falling back to per-page OCR
             (pdf2image + pytesseract) for scanned/image-only PDFs
  - .docx  -> python-docx (paragraphs + tables)
  - .doc   -> best-effort via antiword/textract-style fallback is not bundled;
             we ask the user to re-save as .docx/.pdf if this fails (see note below)
  - .png/.jpg/.jpeg/.webp -> pytesseract OCR directly

Every function returns plain extracted text. Nothing here calls the AI —
this module's only job is turning a file into text so ai_service can read it.
"""
import io
import os
import logging

import pdfplumber
from PIL import Image
import pytesseract
import docx  # python-docx

logger = logging.getLogger("resume_parser")

MIN_NATIVE_TEXT_CHARS_PER_PAGE = 40  # below this, we treat the page as "scanned" and OCR it

# OCR language can be overridden by admin settings; defaults to English.
_OCR_LANG = "eng"


def set_ocr_language(lang: str):
    """Called at startup or when admin changes OCR language settings."""
    global _OCR_LANG
    if lang:
        _OCR_LANG = lang


class ResumeParseError(Exception):
    pass


def _ocr_image(img: "Image.Image") -> str:
    try:
        return pytesseract.image_to_string(img, lang=_OCR_LANG) or ""
    except Exception as exc:  # pytesseract/tesseract not available, corrupt image, etc.
        logger.warning("OCR failed on an image: %s", exc)
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF. Uses the real text layer where present, and
    falls back to OCR (rendering the page to an image and running Tesseract)
    for pages that are scanned images with no selectable text."""
    text_parts = []
    needs_ocr_pages = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = (page.extract_text() or "").strip()
            if len(page_text) >= MIN_NATIVE_TEXT_CHARS_PER_PAGE:
                text_parts.append(page_text)
            else:
                text_parts.append(None)  # placeholder, fill in after OCR pass
                needs_ocr_pages.append(i)

    if needs_ocr_pages:
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(file_bytes, dpi=250)
            for i in needs_ocr_pages:
                if i < len(images):
                    text_parts[i] = _ocr_image(images[i])
        except Exception as exc:
            logger.warning("PDF->image OCR fallback failed: %s", exc)
            for i in needs_ocr_pages:
                if text_parts[i] is None:
                    text_parts[i] = ""

    return "\n".join(p for p in text_parts if p).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text_from_image(file_bytes: bytes) -> str:
    img = Image.open(io.BytesIO(file_bytes))
    return _ocr_image(img).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1").strip()
        except Exception:
            raise ResumeParseError("Could not read the text file. Please re-save it with UTF-8 encoding.")


def extract_text_from_doc(file_bytes: bytes) -> str:
    # Legacy binary .doc has no reliable pure-Python parser in this environment.
    raise ResumeParseError(
        "Legacy .doc files aren't supported for automatic extraction. "
        "Please re-save the resume as .docx or .pdf and upload again."
    )


EXTENSION_HANDLERS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".doc": extract_text_from_doc,
    ".txt": extract_text_from_txt,
    ".png": extract_text_from_image,
    ".jpg": extract_text_from_image,
    ".jpeg": extract_text_from_image,
    ".webp": extract_text_from_image,
}


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    handler = EXTENSION_HANDLERS.get(ext)
    if not handler:
        raise ResumeParseError(
            f"Unsupported file type '{ext}'. Upload a PDF, DOCX, or image (PNG/JPG) resume."
        )
    text = handler(file_bytes)
    if not text or len(text.strip()) < 20:
        raise ResumeParseError(
            "Could not extract readable text from this file. It may be a low-quality "
            "scan, password-protected, or empty — try a clearer copy or paste the text manually."
        )
    return text
